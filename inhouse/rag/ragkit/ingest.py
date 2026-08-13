# -*- coding: utf-8 -*-
"""documents/산출물/ 아래 보고서(md·docx)와 외부공개 PDF ETL 산출물을 읽어 문서
레코드로 변환.
- pdf는 documents/산출물 트리에서는 같은 파일명의 md를 렌더링한 산출물이라 항상
  건너뜀(중복 콘텐츠) — 원본 PDF가 아니라 pdf_extract.py가 만든 .md만 읽는다.
- docx는 python-docx로 문단·표를 텍스트로 펼침(견출 스타일은 #/## 로 보존 → chunk.py가 재사용).

**EXTRA_ROOTS는 "외부공개 가능"으로 확인된 소스만 추가할 것.** 진단모델 전용
(RAG 금지) 자료는 data_lake/semi_structure/pdf_extract/restricted_diagnosis_only/에
따로 있고, 그 경로는 여기서 절대 참조하지 않는다(사용범위 위반 방지 —
documents/0807/메일내용_0807.txt 2번 항목 참고).
"""
from __future__ import annotations

import hashlib
import os
import re
from dataclasses import dataclass
from pathlib import Path

import docx

DATE_RE = re.compile(r"_(\d{6})(?=_|\.|$)")

# 2026-08-11 버그수정: ROOT가 상대경로("documents/산출물")였는데, 표준 실행
# 관례(CLAUDE.md §2: `cd inhouse && python -m rag ...`)로 돌리면 cwd=inhouse/라
# `inhouse/documents/산출물`(존재하지 않음)을 찾게 되고, load_documents()의
# `if not os.path.isdir(r): continue`에 걸려 조용히 스킵됐다 — 이 인덱스가
# 오늘(2026-08-11) 처음 빌드되기 전까지 아무도 눈치채지 못한 채로 있었다(실측
# 확인: 고치기 전엔 문서 4건만 로드됨 — 전부 EXTRA_ROOTS, 본체 61개 md/15개 docx는
# 전부 스킵됐었음). EXTRA_ROOTS는 처음부터 _REPO_ROOT 기준 절대경로였으니 ROOT도
# 같은 패턴으로 맞춘다.
_REPO_ROOT = Path(__file__).resolve().parents[3]
ROOT = str(_REPO_ROOT / "documents/산출물")

# (경로, week에 쓸 태그 접두사) — 각 하위 디렉토리명이 태그 뒤에 붙는다.
# 예: pdf_extract/shareable/komis_해외투자가이드_4개국/ -> week="외부자료:komis_해외투자가이드_4개국"
EXTRA_ROOTS = [
    (_REPO_ROOT / "inhouse/data_lake/semi_structure/pdf_extract/shareable", "외부자료"),
]

# opendataloader-pdf 기본(비-hybrid) 모드는 텍스트 레이어가 없는 스캔형 PDF에서
# 거의 아무 것도 못 뽑는다(2026-08-10 실측: 해외투자가이드 4개국 전부 이 경우 —
# 이미지 태그·표 뼈대만 남고 실제 문장은 0). 그런 파일을 그대로 청킹/임베딩하면
# 빈 벡터만 늘리므로, EXTRA_ROOTS 문서는 최소 실제 텍스트 분량을 통과해야 한다.
_MIN_REAL_CHARS = 300
_IMAGE_TAG_RE = re.compile(r"!\[image \d+\]\([^)]*\)")
_TABLE_NOISE_RE = re.compile(r"<br\s*/?>|\|")


def _real_content_len(md_text: str) -> int:
    """이미지 참조·표 뼈대(빈 셀+<br>+|)를 뺀 실제 문자(글자/숫자, 모든 문자권) 수.
    스캔형 PDF는 표처럼 보이는 빈 칸만 잔뜩 남기고 셀 안이 비어있는 경우가 실제로
    있어(2026-08-10 실측: 해외투자가이드 4개국 전부 이 패턴), 단순 줄길이 합산으론
    안 걸러진다 — <br>/공백만 있는 셀이 라인당 문자수를 부풀림."""
    stripped = _IMAGE_TAG_RE.sub("", md_text)
    stripped = _TABLE_NOISE_RE.sub(" ", stripped)
    return len(re.findall(r"\w", stripped, re.UNICODE))


@dataclass
class DocRecord:
    doc_id: str
    source_path: str
    week: str
    series_key: str
    doc_date: str  # YYMMDD 없으면 ""
    title: str
    ext: str
    raw_text: str


def _doc_id(path: str) -> str:
    return hashlib.md5(path.encode("utf-8")).hexdigest()[:16]


def _series_and_date(basename: str) -> tuple[str, str]:
    """같은 제목이 날짜만 바뀌어 재등장하는 개정판을 묶기 위한 계열 키.
    예: 핵심광물_시스템구성_요약본_260713 / _260716 -> series=핵심광물_시스템구성_요약본"""
    m = list(DATE_RE.finditer(basename))
    if not m:
        return basename, ""
    last = m[-1]
    date = last.group(1)
    series = basename[: last.start()] + basename[last.end():]
    series = series.rstrip("_")
    return series, date


def _extract_docx(path: str) -> str:
    d = docx.Document(path)
    lines = []
    for para in d.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        style = para.style.name if para.style else ""
        if "Heading 1" in style or "Title" in style:
            lines.append("# " + t)
        elif "Heading 2" in style:
            lines.append("## " + t)
        elif "Heading 3" in style:
            lines.append("### " + t)
        else:
            lines.append(t)
    for table in d.tables:
        lines.append("[TABLE]")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append(" | ".join(cells))
        lines.append("[/TABLE]")
    return "\n".join(lines)


def _title_from_text(text: str, fallback: str) -> str:
    for line in text.splitlines():
        line = line.strip()
        if line.startswith("#"):
            return line.lstrip("#").strip()
        if line:
            return line[:120]
    return fallback


def load_documents(root: str = ROOT) -> list[DocRecord]:
    """md는 그대로, docx는 md 짝이 없을 때만(중복 방지) 추출.
    root(documents/산출물) + EXTRA_ROOTS(외부공개 PDF ETL 산출물)를 함께 읽는다."""
    roots = [(root, None)] + [(str(p), tag) for p, tag in EXTRA_ROOTS]

    md_bases: set[str] = set()
    all_files: list[tuple[str, str | None]] = []  # (path, extra_tag)
    for r, tag in roots:
        if not os.path.isdir(r):
            continue
        for dirpath, _dirnames, filenames in os.walk(r):
            for fn in filenames:
                full = os.path.join(dirpath, fn)
                all_files.append((full, tag))
                if fn.lower().endswith(".md") and not fn.lower().endswith(".meta.md"):
                    md_bases.add(os.path.splitext(full)[0])

    docs: list[DocRecord] = []
    for full, tag in sorted(all_files, key=lambda x: x[0]):
        lower = full.lower()
        if lower.endswith(".pdf"):
            continue  # md 렌더링본, 콘텐츠 중복
        if lower.endswith(".meta.md"):
            continue  # 요약 메타는 본문과 별도 취급하지 않음(필요시 후속 확장)
        rel = os.path.relpath(full, ".")
        base_noext, ext = os.path.splitext(full)
        basename = os.path.basename(base_noext)
        if tag is not None:
            # EXTRA_ROOTS: <extra_root>/<source_label>/file.md -> week="태그:source_label"
            rel_to_extra = os.path.relpath(full, next(p for p, t in EXTRA_ROOTS if t == tag))
            source_label = rel_to_extra.split(os.sep)[0]
            week = f"{tag}:{source_label}"
        else:
            parts = rel.split(os.sep)
            week = parts[2] if len(parts) > 2 else ""

        if lower.endswith(".md"):
            with open(full, encoding="utf-8") as f:
                raw = f.read()
        elif lower.endswith(".docx"):
            if base_noext in md_bases:
                continue  # 같은 이름의 md가 있으면 md만 사용(둘 다 있는 경우 W29 2건)
            raw = _extract_docx(full)
        else:
            continue

        if tag is not None and _real_content_len(raw) < _MIN_REAL_CHARS:
            print(f"  [skip 저품질] {rel} — 실제 텍스트 {_real_content_len(raw)}자 미만"
                  f"(OCR/hybrid 모드 필요, 스캔형 PDF로 추정)")
            continue

        series_key, doc_date = _series_and_date(basename)
        title = _title_from_text(raw, basename)
        docs.append(DocRecord(
            doc_id=_doc_id(rel),
            source_path=rel,
            week=week,
            series_key=series_key,
            doc_date=doc_date,
            title=title,
            ext=ext.lstrip("."),
            raw_text=raw,
        ))
    return docs


if __name__ == "__main__":
    docs = load_documents()
    print(f"문서 {len(docs)}건 로드")
    by_series: dict[str, list[str]] = {}
    for d in docs:
        by_series.setdefault(d.series_key, []).append(d.doc_date)
    multi = {k: v for k, v in by_series.items() if len(v) > 1}
    print(f"버전이 2개 이상인 계열: {len(multi)}건")
    for k, v in sorted(multi.items()):
        print(f"  {k}: {sorted(v)}")
