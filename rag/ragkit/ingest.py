# -*- coding: utf-8 -*-
"""documents/산출물/ 아래 보고서(md·docx)를 읽어 문서 레코드로 변환.
- pdf는 같은 파일명의 md를 렌더링한 산출물이라 항상 건너뜀(중복 콘텐츠).
- docx는 python-docx로 문단·표를 텍스트로 펼침(견출 스타일은 #/## 로 보존 → chunk.py가 재사용).
"""
from __future__ import annotations

import hashlib
import os
import re
from dataclasses import dataclass

import docx

ROOT = "documents/산출물"
DATE_RE = re.compile(r"_(\d{6})(?=_|\.|$)")


@dataclass
class DocRecord:
    doc_id: str
    source_path: str
    week: str
    series_key: str
    doc_date: str  # YYMMDD 없으면 ""
    title: str
    ext: str
    raw_text: str


def _doc_id(path: str) -> str:
    return hashlib.md5(path.encode("utf-8")).hexdigest()[:16]


def _series_and_date(basename: str) -> tuple[str, str]:
    """같은 제목이 날짜만 바뀌어 재등장하는 개정판을 묶기 위한 계열 키.
    예: 핵심광물_시스템구성_요약본_260713 / _260716 -> series=핵심광물_시스템구성_요약본"""
    m = list(DATE_RE.finditer(basename))
    if not m:
        return basename, ""
    last = m[-1]
    date = last.group(1)
    series = basename[: last.start()] + basename[last.end():]
    series = series.rstrip("_")
    return series, date


def _extract_docx(path: str) -> str:
    d = docx.Document(path)
    lines = []
    for para in d.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        style = para.style.name if para.style else ""
        if "Heading 1" in style or "Title" in style:
            lines.append("# " + t)
        elif "Heading 2" in style:
            lines.append("## " + t)
        elif "Heading 3" in style:
            lines.append("### " + t)
        else:
            lines.append(t)
    for table in d.tables:
        lines.append("[TABLE]")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append(" | ".join(cells))
        lines.append("[/TABLE]")
    return "\n".join(lines)


def _title_from_text(text: str, fallback: str) -> str:
    for line in text.splitlines():
        line = line.strip()
        if line.startswith("#"):
            return line.lstrip("#").strip()
        if line:
            return line[:120]
    return fallback


def load_documents(root: str = ROOT) -> list[DocRecord]:
    """md는 그대로, docx는 md 짝이 없을 때만(중복 방지) 추출."""
    md_bases: set[str] = set()
    all_files: list[str] = []
    for dirpath, _dirnames, filenames in os.walk(root):
        for fn in filenames:
            full = os.path.join(dirpath, fn)
            all_files.append(full)
            if fn.lower().endswith(".md") and not fn.lower().endswith(".meta.md"):
                md_bases.add(os.path.splitext(full)[0])

    docs: list[DocRecord] = []
    for full in sorted(all_files):
        lower = full.lower()
        if lower.endswith(".pdf"):
            continue  # md 렌더링본, 콘텐츠 중복
        if lower.endswith(".meta.md"):
            continue  # 요약 메타는 본문과 별도 취급하지 않음(필요시 후속 확장)
        rel = os.path.relpath(full, ".")
        parts = rel.split(os.sep)
        week = parts[2] if len(parts) > 2 else ""
        base_noext, ext = os.path.splitext(full)
        basename = os.path.basename(base_noext)

        if lower.endswith(".md"):
            with open(full, encoding="utf-8") as f:
                raw = f.read()
        elif lower.endswith(".docx"):
            if base_noext in md_bases:
                continue  # 같은 이름의 md가 있으면 md만 사용(둘 다 있는 경우 W29 2건)
            raw = _extract_docx(full)
        else:
            continue

        series_key, doc_date = _series_and_date(basename)
        title = _title_from_text(raw, basename)
        docs.append(DocRecord(
            doc_id=_doc_id(rel),
            source_path=rel,
            week=week,
            series_key=series_key,
            doc_date=doc_date,
            title=title,
            ext=ext.lstrip("."),
            raw_text=raw,
        ))
    return docs


if __name__ == "__main__":
    docs = load_documents()
    print(f"문서 {len(docs)}건 로드")
    by_series: dict[str, list[str]] = {}
    for d in docs:
        by_series.setdefault(d.series_key, []).append(d.doc_date)
    multi = {k: v for k, v in by_series.items() if len(v) > 1}
    print(f"버전이 2개 이상인 계열: {len(multi)}건")
    for k, v in sorted(multi.items()):
        print(f"  {k}: {sorted(v)}")
